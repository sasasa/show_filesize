import PySimpleGUI as sg
import os
import re
from pathlib import Path
from chardet import detect
from openpyxl import load_workbook
from docx import Document
from pdfminer.high_level import extract_text

#【2.アプリに表示する文字列を設定】
title = "ファイルの合計サイズを表示（フォルダ以下すべての）"
infolder = "."
label1, value1 = "拡張子", "*"

# 配列をflattenする
def flatten(x):
    return [z for y in x for z in (flatten(y) if hasattr(y, '__iter__') and not isinstance(y, str) else (y,))]

#【3.関数: ファイルサイズを最適単位で返す】
def format_bytes(size):
    units = ["バイト","KB","MB","GB","TB","PB","EB"]
    n = 0
    while size > 1024:
        size = size / 1024.0
        n += 1
    return str(int(size)) + " " + units[n]

def copyFiles():
    global copyList, filelist, grepList, grepCopyList
    if len(grepCopyList) > 0:
        for i, path in enumerate(grepCopyList):
            # フォルダが存在しなければ作成する
            if not os.path.exists(os.path.dirname(path)):
                os.makedirs(os.path.dirname(path))
            # ファイルをコピーする
            with open(grepList[i], "rb") as f:
                with open(path, "wb") as f2:
                    f2.write(f.read())
        return
    if len(copyList) > 0:
        for i, path in enumerate(copyList):
            # フォルダが存在しなければ作成する
            if not os.path.exists(os.path.dirname(path)):
                os.makedirs(os.path.dirname(path))
            # ファイルをコピーする
            with open(filelist[i], "rb") as f:
                with open(path, "wb") as f2:
                    f2.write(f.read())

#【3.関数: フォルダ以下のファイルのサイズ合計を求める】
def foldersize(infolder, ext, extList, search):
    global itms, copyList, filelist, grepList, grepCopyList
    msg = ""
    allsize = 0
    filelist = []
    copyList = []
    grepList = []
    grepCopyList = []
    extSet = set()
    grepExtSet = set()
    # "全て"を含むか
    if len(extList) > 0:
        # 拡張子配列のそれぞれの.を削除
        ext = [x.replace(".", "") for x in extList]
        if "全て" in extList:
            ext = ["*"]
    else:
        ext = [ext]
    try:
        if values["radio1"]:
            searchFiles = getattr(Path(infolder), 'glob')#このフォルダのみのファイルを
        elif values["radio2"]:
            searchFiles = getattr(Path(infolder), 'rglob')#このフォルダ以下すべてのファイルを
        for p in flatten([searchFiles(f"*.{x}") for x in ext]):
            if p.name and p.name[0] != "." and p.is_file():                #隠しファイルでなければ
                # コピー先パスを取得する
                # パスが存在しているならば
                if values["outfolder"] != "" and values["outfolder"] != infolder and values["outfolder"] != "." and values["outfolder"] != "./" and os.path.isdir(values["outfolder"]):
                    # コピーするファイルのパスを取得する
                    copyList.append((values["outfolder"] + "/" + os.path.relpath(p, infolder)))
                filelist.append(str(p))         #リストに追加して
        for filename in sorted(filelist):       #ソートして1ファイルずつ処理
            # 拡張子を取得 小文字に変換
            ext = os.path.splitext(filename)[1].lower()
            extSet.add(ext)
            path = Path(filename)
            if search != "":
                # search = re.compile(repr(search)[1:-1])
                search = re.compile(search)
                # エクセルファイルかどうか
                if ext == ".xlsx":
                    # エクセルファイルを開く
                    wb = load_workbook(filename)
                    # シート名を取得
                    flg = True
                    for sheetname in wb.sheetnames:
                        if flg == False:
                            break
                        # シートを選択
                        sheet = wb[sheetname]
                        # セルを取得
                        for row in sheet.rows:
                            if flg == False:
                                break
                            for cell in row:
                                if flg == False:
                                    break
                                # 正規表現で検索文字が含まれているか
                                if len(re.findall(search, str(cell.value))) > 0:
                                    # ファイルサイズを取得
                                    size = path.stat().st_size
                                    msg += filename + " : "+format_bytes(size)+"\n"
                                    allsize += size
                                    grepList.append(filename)
                                    if values["outfolder"] != "" and values["outfolder"] != infolder and values["outfolder"] != "." and values["outfolder"] != "./" and os.path.isdir(values["outfolder"]):
                                        # コピーするファイルのパスを取得する
                                        grepCopyList.append((values["outfolder"] + "/" + os.path.relpath(path, infolder)))
                                    grepExtSet.add(ext)
                                    flg = False
                                    break
                # ワードファイルかどうか
                elif ext == ".doc" or ext == ".docx":
                    # ワードファイルを開く
                    try:
                        doc = Document(filename)
                    except:
                        continue
                    # テキストを取得
                    flg = True
                    for paragraph in doc.paragraphs:
                        # 検索文字が含まれているか
                        if len(re.findall(search, paragraph.text)) > 0:
                            # ファイルサイズを取得
                            size = path.stat().st_size
                            msg += filename + " : "+format_bytes(size)+"\n"
                            allsize += size
                            grepList.append(filename)
                            if values["outfolder"] != "" and values["outfolder"] != infolder and values["outfolder"] != "." and values["outfolder"] != "./" and os.path.isdir(values["outfolder"]):
                                # コピーするファイルのパスを取得する
                                grepCopyList.append((values["outfolder"] + "/" + os.path.relpath(path, infolder)))
                            grepExtSet.add(ext)
                            flg = False
                            break
                    # テーブルを取得
                    for t in doc.tables:
                        if flg == False:
                            break
                        for row in t.rows:
                            if flg == False:
                                break
                            for cell in row.cells:
                                # 検索文字が含まれているか
                                if len(re.findall(search, cell.text)) > 0:
                                    # ファイルサイズを取得
                                    size = path.stat().st_size
                                    msg += filename + " : "+format_bytes(size)+"\n"
                                    allsize += size
                                    grepList.append(filename)
                                    if values["outfolder"] != "" and values["outfolder"] != infolder and values["outfolder"] != "." and values["outfolder"] != "./" and os.path.isdir(values["outfolder"]):
                                        # コピーするファイルのパスを取得する
                                        grepCopyList.append((values["outfolder"] + "/" + os.path.relpath(path, infolder)))
                                    grepExtSet.add(ext)
                                    flg = False
                                    break
                # PDFファイルかどうか
                elif ext == ".pdf":
                    try:
                        txt = extract_text(filename)
                        # 検索文字が含まれているか
                        if len(re.findall(search, txt)) > 0:
                            # ファイルサイズを取得
                            size = path.stat().st_size
                            msg += filename + " : "+format_bytes(size)+"\n"
                            allsize += size
                            grepList.append(filename)
                            if values["outfolder"] != "" and values["outfolder"] != infolder and values["outfolder"] != "." and values["outfolder"] != "./" and os.path.isdir(values["outfolder"]):
                                # コピーするファイルのパスを取得する
                                grepCopyList.append((values["outfolder"] + "/" + os.path.relpath(path, infolder)))
                            grepExtSet.add(ext)
                    except:
                        pass
                # テキストファイルかどうか
                else:
                    with open(path, "rb") as f:
                        b = f.read(4096)
                        if b:
                            try:
                                encode = detect(b)["encoding"]
                                # テキストファイルかどうか
                                if encode != None:
                                    if encode == "Windows-1252":
                                        encode = "SHIFT-JIS"
                                    txt = path.read_text(encoding=encode)
                                    if len(re.findall(search, txt)) > 0:
                                        # ファイルサイズを取得
                                        size = path.stat().st_size
                                        msg += filename + " : "+format_bytes(size)+"\n"
                                        allsize += size
                                        grepList.append(filename)
                                        if values["outfolder"] != "" and values["outfolder"] != infolder and values["outfolder"] != "." and values["outfolder"] != "./" and os.path.isdir(values["outfolder"]):
                                            # コピーするファイルのパスを取得する
                                            grepCopyList.append((values["outfolder"] + "/" + os.path.relpath(path, infolder)))
                                        grepExtSet.add(ext)
                            except:
                                pass
            else:
                # ファイルサイズを取得
                size = path.stat().st_size
                msg += filename + " : "+format_bytes(size)+"\n"
                allsize += size
        filesize = "合計サイズ = " + format_bytes(allsize) + "\n"
        if search != "":
            filesize += "ファイル数 = " + str(len(grepList))+ "\n"
        else:
            filesize += "ファイル数 = " + str(len(filelist))+ "\n"
        msg = filesize + msg
        # リストボックスの更新 ソートする
        if search != "":
            extSet = grepExtSet
        itms = itms + sorted(list(extSet))
        window["listbox1"].update(itms)
        # リストボックスのサイズを更新
        if len(itms) // 2 > 7:
            height = 7
        else:
            if 2 <= len(itms) <= 4:
                height = 2
            else:
                height = len(itms) // 2
        window["listbox1"].Widget.config(height=height)
        # リストボックスの選択状態を更新
        # 拡張子配列のそれぞれのindexを取得 itms配列の中にextList配列が含まれているか
        window["listbox1"].update(set_to_index=[itms.index(x) for x in extList if x in itms])
        itms = ["全て",]
        return msg
    except Exception as e:
        sg.popup(e, title='エラー', keep_on_top=True)
        return
#--------------------^^^
def execute():
    infolder = values["infolder"]
    outfolder = values["outfolder"]
    if infolder == outfolder:
        sg.popup("読み込みフォルダと書き込みフォルダが同じです", title='エラー', keep_on_top=True)
        return
    value1 = values["input1"]
    search = values["input2"]
    # 絞り込む拡張子を取得
    extList = values["listbox1"]
    if value1 == "":
        # popupを出す
        sg.popup("拡張子を入力してください", title='エラー', keep_on_top=True)
        return
    msg = foldersize(infolder, value1, extList, search)
    window["text1"].update(msg)
#アプリのレイアウト
itms = ["全て",]
layout = [[sg.Text("読み込みフォルダ", size=(14,1)),
           sg.Input(infolder, key="infolder"),sg.FolderBrowse("選択")],
           [sg.Text("書き込みフォルダ", size=(14,1)),
           sg.Input("", key="outfolder"),sg.FolderBrowse("選択")],
          [sg.Text(label1, size=(14,1)), sg.Input(value1, key="input1")],
          [sg.Text("検索文字", size=(14,1)), sg.Input("", key="input2")],
          [sg.Button("検索", size=(20,1), pad=(5,5), bind_return_key=True)],
          [sg.Button("フォルダごとコピー", size=(20,1), bind_return_key=True)],
          [sg.Radio('このフォルダだけ', group_id="RADIO", default=True, key="radio1", enable_events=True),
           sg.Radio('このフォルダ配下全て', group_id="RADIO", default=False, key="radio2", enable_events=True),],
          [sg.Listbox(itms, size=(35,len(itms)), default_values=[["全て"]], key="listbox1", enable_events=True, select_mode=sg.LISTBOX_SELECT_MODE_MULTIPLE),],
          [sg.Multiline(key="text1", size=(60,10))]]
# Listboxを選択状態にする
#アプリの実行処理
window = sg.Window(title, layout, font=(None,14), resizable=True, finalize=True)
window["listbox1"].update(set_to_index=[0])
while True:
    event, values = window.read()
    # print(" イベント:",event ,", 値:",values)
    if event == None:
        break
    if event == "検索":
        execute()
    if event == "フォルダごとコピー":
        execute()
        copyFiles()
    if event == "radio1":
        execute()
    if event == "radio2":
        execute()
    if event == "listbox1":
        execute()
window.close()
